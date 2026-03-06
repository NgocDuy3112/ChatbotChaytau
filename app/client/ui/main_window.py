from __future__ import annotations

import html
import csv
import re
import uuid
from datetime import datetime
from pathlib import Path

from PyQt6.QtCore import QEvent, QObject, QSize, QSettings, Qt, QTimer, QUrl
from PyQt6.QtGui import QColor, QDesktopServices, QIcon, QPainter, QTextDocument
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtWidgets import (
    QCheckBox,
    QComboBox,
    QFileDialog,
    QHBoxLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QMenu,
    QPushButton,
    QScrollArea,
    QSplitter,
    QStatusBar,
    QTextBrowser,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from ..api.client import ApiError, ChatApiClient
from ..models.dto import BaseMessage, ChatRequest, Conversation
from ..state.store import ChatMessage, ChatState
from ..workers.stream_worker import ChatStreamWorker, StreamResult
from ..utils.resources import get_icons_dir, get_instructions_dir, get_sheets_dir


class WheelEventFilter(QObject):
    def eventFilter(self, obj: QObject, event: QEvent) -> bool:
        if event.type() == QEvent.Type.Wheel:
            if isinstance(obj, QComboBox):
                if not obj.view().isVisible():
                    event.ignore()
                    return True
                return super().eventFilter(obj, event)

            if isinstance(obj, QLineEdit) and isinstance(obj.parent(), QComboBox):
                combo = obj.parent()
                if isinstance(combo, QComboBox) and not combo.view().isVisible():
                    event.ignore()
                    return True
                return super().eventFilter(obj, event)

            # Nếu widget chưa được focus thì ignore wheel event để nó tự trả về cho ScrollArea cha
            if not obj.hasFocus():
                event.ignore()
                return True
        return super().eventFilter(obj, event)


class MainWindow(QMainWindow):
    _WINDOW_SIZE = (1500, 800)
    _MAIN_SPLITTER_SIZES = [240, 960]
    _RIGHT_SPLITTER_SIZES = [520, 300]
    _INPUT_MIN_HEIGHT = 42
    _INPUT_MAX_HEIGHT = 128
    _IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}
    _PROMPT_COMPANY_KEY = "ten_cong_ty"
    _PROMPT_ROLE_KEY = "vai_tro"
    _PROMPT_WORK_KEY = "noi_dung_chi_tiet"
    _HIDDEN_PROMPT_KEYS = {"muc_tieu", "noi_dung_chi_tiet"}
    _BASE_INSTRUCTION_PREFIXES = ("0_", "00_")
    _DEFAULT_OVERLAY_PREFIXES = ("3_", "03_")
    _TEMPLATE_INSTRUCTION_PREFIXES = ("5_", "01_")
    _WORK_INSTRUCTION_PREFIXES = ("1_", "2_", "3_", "4_")
    _MODEL_OPTIONS = (
        "gemini-3.1-flash-lite-preview",
        "gemini-3-flash-preview",
    )

    def __init__(self, base_url: str | None = None):
        super().__init__()
        self.setWindowTitle("Chatbot Desktop")
        self.setFixedSize(*self._WINDOW_SIZE)

        self.client = ChatApiClient(base_url=base_url or "http://localhost:8000")
        self.state = ChatState()
        self.stream_worker: ChatStreamWorker | None = None
        self.settings = QSettings("ChatbotChaytau", "ChatbotDesktop")
        self.fixed_model = self._MODEL_OPTIONS[0]
        self.model_selector: QComboBox | None = None
        self.prompt_template_text = ""
        self.prompt_field_inputs: dict[str, QWidget] = {}
        self.prompt_options: dict[str, list[str]] = {}
        self.work_prompt_map: dict[str, str] = {}
        self.company_context_by_name: dict[str, str] = {}
        self.company_context_lookup: dict[str, str] = {}
        self.company_context_checkbox: QCheckBox | None = None
        self.search_grounding_checkbox: QCheckBox | None = None
        self.auto_open_export_checkbox: QCheckBox | None = None
        self._response_status_text = "Trạng thái phản hồi: Sẵn sàng"
        self._response_status_state = "idle"
        self._spinner_frames = ("⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏")
        self._spinner_index = 0
        self._spinner_timer = QTimer(self)
        self._spinner_timer.setInterval(90)
        self._spinner_timer.timeout.connect(self._advance_response_spinner)
        self._restoring_right_panel_settings = False
        self.default_instruction_profile_text = self._load_default_instruction_profile_text()
        self.default_instructions_text = self.default_instruction_profile_text
        self.wheel_event_filter = WheelEventFilter(self)
        (
            self.prompt_options,
            self.company_context_by_name,
            self.company_context_lookup,
            self.work_prompt_map,
        ) = self._load_prompt_bundle_data()

        self._build_ui()
        self._apply_styles()
        self._load_settings()
        self._load_conversations()

    def _build_ui(self) -> None:
        root = QWidget(self)
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(14, 14, 14, 12)
        root_layout.setSpacing(10)

        top_bar = QHBoxLayout()
        top_bar.setSpacing(8)

        title_label = QLabel("Trợ lý trò chuyện")
        title_label.setObjectName("appTitle")
        top_bar.addWidget(title_label)
        top_bar.addStretch(1)
        root_layout.addLayout(top_bar)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(1)

        left_panel = QWidget()
        left_panel.setObjectName("leftPanel")
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(8, 8, 8, 8)
        left_layout.setSpacing(8)

        section_label = QLabel("Cuộc trò chuyện")
        section_label.setObjectName("sectionLabel")
        left_layout.addWidget(section_label)

        new_chat_btn = QPushButton("+ Trò chuyện mới")
        new_chat_btn.setObjectName("newChatButton")
        new_chat_btn.clicked.connect(self._new_chat)
        left_layout.addWidget(new_chat_btn)

        self.conversation_list = QListWidget()
        self.conversation_list.setObjectName("conversationList")
        self.conversation_list.itemSelectionChanged.connect(self._on_conversation_selected)
        self.conversation_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.conversation_list.customContextMenuRequested.connect(self._on_conversation_context_menu)
        left_layout.addWidget(self.conversation_list)
        splitter.addWidget(left_panel)

        right_panel = QWidget()
        right_panel.setObjectName("rightPanel")
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(8, 8, 8, 8)
        right_layout.setSpacing(8)

        # Keep setting checkboxes near composer action buttons (not in top row)
        self.search_grounding_checkbox = QCheckBox("Tìm kiếm bằng Google")
        self.search_grounding_checkbox.setChecked(True)
        self.search_grounding_checkbox.toggled.connect(self._on_right_panel_setting_changed)

        self.auto_open_export_checkbox = QCheckBox("Tự mở file sau khi xuất")
        self.auto_open_export_checkbox.setChecked(True)
        self.auto_open_export_checkbox.toggled.connect(self._on_right_panel_setting_changed)

        right_splitter = QSplitter(Qt.Orientation.Horizontal)
        right_splitter.setHandleWidth(1)

        chat_container = QWidget()
        chat_layout = QVBoxLayout(chat_container)
        chat_layout.setContentsMargins(0, 0, 0, 0)
        chat_layout.setSpacing(8)

        self.chat_view = QTextBrowser()
        self.chat_view.setObjectName("chatView")
        self.chat_view.setOpenLinks(False)
        self.chat_view.anchorClicked.connect(self._on_chat_link_clicked)
        chat_layout.addWidget(self.chat_view, 1)

        self.add_file_button = QPushButton()
        self.add_file_button.setObjectName("addFileButton")
        self.add_file_button.setFixedSize(42, 42)
        self.add_file_button.clicked.connect(self._attach_files)
        self._configure_icon_button(
            self.add_file_button,
            ("mail-attachment", "attachment"),
            "📎",
            "Thêm file",
            icon_size=18,
        )

        self.add_image_button = QPushButton()
        self.add_image_button.setObjectName("addImageButton")
        self.add_image_button.setFixedSize(42, 42)
        self.add_image_button.clicked.connect(self._attach_images)
        self._configure_resource_svg_icon_button(
            self.add_image_button,
            "image.svg",
            ("image-x-generic", "insert-image"),
            "🖼",
            "Thêm ảnh",
            icon_size=18,
        )

        self.clear_file_button = QPushButton()
        self.clear_file_button.setObjectName("clearFileButton")
        self.clear_file_button.setFixedSize(42, 42)
        self.clear_file_button.clicked.connect(self._clear_attachments)
        self._configure_icon_button(
            self.clear_file_button,
            ("user-trash", "edit-delete"),
            "🗑",
            "Xóa file",
            icon_size=18,
        )

        self.send_button = QPushButton()
        self.send_button.setObjectName("sendButton")
        self.send_button.setFixedSize(42, 42)
        self.send_button.clicked.connect(self._send_message)
        self._configure_icon_button(
            self.send_button,
            ("mail-send", "document-send"),
            "➤",
            "Gửi",
            icon_size=18,
        )

        self.model_selector = QComboBox()
        self.model_selector.setObjectName("modelSelector")
        self.model_selector.addItems(list(self._MODEL_OPTIONS))
        self.model_selector.setCurrentText(self.fixed_model)
        self.model_selector.setToolTip("Chọn model phản hồi")
        self.model_selector.currentTextChanged.connect(self._on_right_panel_setting_changed)

        attachments_row = QHBoxLayout()
        attachments_row.setSpacing(8)

        self.attachment_list = QTextBrowser()
        self.attachment_list.setObjectName("attachmentList")
        self.attachment_list.setOpenLinks(False)
        self.attachment_list.setOpenExternalLinks(False)
        self.attachment_list.setMinimumHeight(112)
        self.attachment_list.setMaximumHeight(112)

        attachments_row.addWidget(self.attachment_list, 1)
        chat_layout.addLayout(attachments_row)

        input_panel = QWidget()
        input_panel.setObjectName("inputComposer")
        input_panel_layout = QVBoxLayout(input_panel)
        input_panel_layout.setContentsMargins(10, 8, 10, 8)
        input_panel_layout.setSpacing(6)

        self.input_box = QTextEdit()
        self.input_box.setObjectName("inputBox")
        self.input_box.setPlaceholderText("Nhập tin nhắn...")
        self.input_box.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.input_box.installEventFilter(self)
        self.input_box.textChanged.connect(self._adjust_input_box_height)
        input_panel_layout.addWidget(self.input_box)

        actions_row = QHBoxLayout()
        actions_row.setContentsMargins(0, 0, 0, 0)
        actions_row.setSpacing(8)
        actions_row.addWidget(self.add_file_button)
        actions_row.addWidget(self.add_image_button)
        actions_row.addWidget(self.clear_file_button)
        actions_row.addWidget(self.search_grounding_checkbox)
        actions_row.addWidget(self.auto_open_export_checkbox)
        actions_row.addStretch(1)
        actions_row.addWidget(self.model_selector)
        actions_row.addWidget(self.send_button)
        input_panel_layout.addLayout(actions_row)

        chat_layout.addWidget(input_panel)

        self._adjust_input_box_height()

        # Response status and spinner are rendered inside the assistant message bubble
        # (no global status row here anymore).

        self._update_attachment_label()
        right_splitter.addWidget(chat_container)

        prompt_sidebar = QWidget()
        prompt_sidebar.setObjectName("promptSidebar")
        prompt_layout = QVBoxLayout(prompt_sidebar)
        prompt_layout.setContentsMargins(8, 8, 8, 8)
        prompt_layout.setSpacing(8)

        prompt_scroll = QScrollArea()
        prompt_scroll.setObjectName("promptScroll")
        prompt_scroll.setWidgetResizable(True)
        prompt_form_widget = QWidget()
        self.prompt_form_layout = QVBoxLayout(prompt_form_widget)
        self.prompt_form_layout.setContentsMargins(0, 0, 0, 0)
        self.prompt_form_layout.setSpacing(6)
        prompt_scroll.setWidget(prompt_form_widget)
        prompt_layout.addWidget(prompt_scroll, 1)

        self._load_prompt_sidebar_fields()

        right_splitter.addWidget(prompt_sidebar)
        right_splitter.setSizes(self._RIGHT_SPLITTER_SIZES)
        right_layout.addWidget(right_splitter, 1)

        splitter.addWidget(right_panel)
        splitter.setSizes(self._MAIN_SPLITTER_SIZES)
        root_layout.addWidget(splitter, 1)

        self.setCentralWidget(root)
        self.setStatusBar(QStatusBar())
        self.statusBar().showMessage("Sẵn sàng")

    def eventFilter(self, watched: QObject, event: QEvent) -> bool:
        if watched is self.input_box and event.type() == QEvent.Type.KeyPress:
            key = getattr(event, "key", lambda: None)()
            if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                modifiers = getattr(event, "modifiers", lambda: Qt.KeyboardModifier.NoModifier)()
                if modifiers & Qt.KeyboardModifier.ShiftModifier:
                    return False

                if self.send_button.isEnabled():
                    self._send_message()
                return True

        return super().eventFilter(watched, event)

    def _apply_styles(self) -> None:
        self.setStyleSheet(
            """
            QWidget {
                font-size: 13px;
            }

            QMainWindow {
                background: #f4f5f7;
            }

            #appTitle {
                font-size: 16px;
                font-weight: 700;
                color: #1f2937;
            }

            #leftPanel, #rightPanel {
                background: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 10px;
            }

            #sectionLabel {
                font-size: 12px;
                font-weight: 600;
                color: #6b7280;
                padding-left: 2px;
            }

            #fieldLabel {
                color: #4b5563;
                font-weight: 600;
            }

            QPushButton {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                background: #ffffff;
                padding: 6px 12px;
            }

            QPushButton:hover {
                background: #f3f4f6;
            }

            QPushButton:disabled {
                color: #9ca3af;
                background: #f3f4f6;
            }

            #newChatButton {
                background: #2563eb;
                color: white;
                border: 1px solid #2563eb;
                font-weight: 600;
            }

            #newChatButton:hover {
                background: #1d4ed8;
            }

            #addFileButton {
                background: #2563eb;
                color: #ffffff;
                border: 1px solid #2563eb;
                font-weight: 600;
                padding: 0;
            }

            #addFileButton:hover {
                background: #1d4ed8;
            }

            #addFileButton:disabled {
                background: #93c5fd;
                border: 1px solid #93c5fd;
                color: rgba(255, 255, 255, 0.55);
            }

            #addImageButton {
                background: #7c3aed;
                color: #ffffff;
                border: 1px solid #7c3aed;
                font-weight: 600;
                padding: 0;
            }

            #addImageButton:hover {
                background: #6d28d9;
            }

            #addImageButton:disabled {
                background: #c4b5fd;
                border: 1px solid #c4b5fd;
                color: rgba(255, 255, 255, 0.55);
            }

            #clearFileButton {
                background: #dc2626;
                color: #ffffff;
                border: 1px solid #dc2626;
                font-weight: 600;
                padding: 0;
            }

            #clearFileButton:hover {
                background: #b91c1c;
            }

            #clearFileButton:disabled {
                background: #fca5a5;
                border: 1px solid #fca5a5;
                color: rgba(255, 255, 255, 0.55);
            }

            #sendButton {
                background: #16a34a;
                color: #ffffff;
                border: 1px solid #16a34a;
                font-weight: 600;
                padding: 0;
            }

            #sendButton:hover {
                background: #15803d;
            }

            #sendButton:disabled {
                background: #86efac;
                border: 1px solid #86efac;
                color: rgba(255, 255, 255, 0.55);
            }

            #conversationList {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                outline: 0;
                padding: 4px;
            }

            #conversationList::item {
                padding: 10px 8px;
                border-radius: 6px;
            }

            #conversationList::item:selected {
                background: #e5edff;
                color: #1e3a8a;
            }

            #chatView {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background: #fafafa;
                padding: 8px;
            }

            #instructionsInput, #modelInput {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                background: #ffffff;
                padding: 6px;
            }

            #modelSelector {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                background: #ffffff;
                padding: 4px 8px;
                min-width: 200px;
                min-height: 32px;
            }

            #modelSelector:disabled {
                color: #9ca3af;
                background: #f3f4f6;
            }

            #inputComposer {
                border: 1px solid #d1d5db;
                border-radius: 14px;
                background: #ffffff;
            }

            #inputBox {
                border: 0;
                background: transparent;
                padding: 2px 2px 0 2px;
            }

            #promptSidebar {
                background: #f8fafc;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
            }

            #promptHint {
                color: #6b7280;
                font-size: 12px;
                margin-bottom: 4px;
            }

            /* Style cho input và combo trong sidebar */
            #promptFieldInput {
                border: 1px solid #d1d5db;
                border-radius: 6px;
                background: #ffffff;
                padding: 4px 8px;
                min-height: 28px;
            }

            QComboBox QAbstractItemView {
                min-width: 200px;
            }

            /* Đảm bảo QComboBox có mũi tên rõ ràng */
            QComboBox#promptFieldInput {
                padding: 4px 8px; 
                border: 1px solid #d1d5db;
                border-radius: 6px;
                background: #ffffff;
                min-height: 28px;
            }

            QComboBox#promptFieldInput::drop-down {
                subcontrol-origin: border;
                subcontrol-position: top right;
                width: 30px;
                border-left: 1px solid #d1d5db;
                border-top-right-radius: 6px;
                border-bottom-right-radius: 6px;
                background: #f3f4f6;
            }

            QComboBox#promptFieldInput::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid #4b5563;
                width: 0;
                height: 0;
            }

            QComboBox#promptFieldInput::drop-down:hover {
                background: #e5e7eb;
            }

            QComboBox#promptFieldInput:on {
                border-bottom-left-radius: 0px;
                border-bottom-right-radius: 0px;
            }

            QComboBox#promptFieldInput:hover, QTextEdit#promptFieldInput:hover {
                border-color: #3b82f6;
            }

            QTextEdit#promptFieldInput {
                padding: 6px;
            }

            #attachmentList {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                background: #f9fafb;
                padding: 4px;
                color: #4b5563;
            }

            /* Bong bóng Chat */
            .message-container {
                margin: 10px 0;
            }

            .user-bubble {
                background-color: #007bff;
                color: white;
                border-radius: 15px 15px 0 15px;
                padding: 10px 15px;
                margin-left: 50px;
                margin-right: 5px;
                box-shadow: 0 1px 2px rgba(0,0,0,0.1);
            }

            .assistant-bubble {
                background-color: #f1f3f4;
                color: #202124;
                border-radius: 15px 15px 15px 0;
                padding: 12px 18px;
                margin-right: 50px;
                margin-left: 5px;
                border: 1px solid #e0e0e0;
            }

            .message-meta {
                font-size: 11px;
                color: #70757a;
                margin-bottom: 4px;
            }

            /* Markdown Enhancements */
            code {
                background-color: rgba(0,0,0,0.05);
                padding: 2px 4px;
                border-radius: 4px;
                font-family: \"Consolas\", \"Monaco\", monospace;
            }

            pre {
                background-color: #2d2d2d;
                color: #f8f8f2;
                padding: 12px;
                border-radius: 8px;
                margin: 10px 0;
                font-family: \"Consolas\", \"Monaco\", monospace;
            }

            table {
                border-collapse: collapse;
                width: 100%;
                margin: 10px 0;
            }

            th, td {
                border: 1px solid #dfe1e5;
                padding: 8px;
                text-align: left;
            }

            th {
                background-color: #f8f9fa;
            }
            """
        )

    def _new_chat(self) -> None:
        self.state.reset_chat()
        self.conversation_list.clearSelection()
        self._render_messages()
        self._update_attachment_label()
        self._set_response_status("Trạng thái phản hồi: Sẵn sàng", "idle")
        self.statusBar().showMessage("Đã tạo cuộc trò chuyện mới", 3000)

    def _load_settings(self) -> None:
        self._restoring_right_panel_settings = True
        try:
            instructions_raw = self.settings.value("chat/default_instructions", "")
            search_grounding_raw = self.settings.value("chat/search_grounding_enabled", True)
            auto_open_exports_raw = self.settings.value("export/auto_open_exported_files", True)
            selected_model_raw = self.settings.value("chat/model_name", self.fixed_model)

            instructions = str(instructions_raw or "").strip()
            self.default_instructions_text = instructions or self.default_instruction_profile_text

            if self.search_grounding_checkbox is not None:
                self.search_grounding_checkbox.setChecked(self._coerce_setting_bool(search_grounding_raw))

            if self.auto_open_export_checkbox is not None:
                self.auto_open_export_checkbox.setChecked(self._coerce_setting_bool(auto_open_exports_raw))

            selected_model = str(selected_model_raw or self.fixed_model).strip()
            if selected_model not in self._MODEL_OPTIONS:
                selected_model = self.fixed_model

            if self.model_selector is not None:
                model_index = self.model_selector.findText(selected_model)
                self.model_selector.setCurrentIndex(model_index if model_index >= 0 else 0)

            self.fixed_model = selected_model
        finally:
            self._restoring_right_panel_settings = False

    def _load_prompt_sidebar_fields(self) -> None:
        self.prompt_template_text = self._load_prompt_template_text()
        self.prompt_field_inputs.clear()
        self.company_context_checkbox = None

        self._clear_prompt_form_layout()

        if not self.prompt_template_text:
            self._add_prompt_sidebar_fallback(
                "Không tìm thấy file template prompt trong resources/instructions."
            )
            return

        placeholders = self._extract_unique_placeholders(self.prompt_template_text)
        if not placeholders:
            self._add_prompt_sidebar_fallback(
                "Prompt 01 không có placeholder dạng {{variable}} để tạo input field."
            )
            return

        var_config = self._prompt_variable_config()

        for placeholder in placeholders:
            if placeholder in self._HIDDEN_PROMPT_KEYS:
                continue
            self._add_prompt_field(placeholder, var_config.get(placeholder, {}))

        self._add_company_context_checkbox_if_available()

        self._restore_prompt_sidebar_settings()
        self.prompt_form_layout.addStretch(1)

    def _clear_prompt_form_layout(self) -> None:
        while self.prompt_form_layout.count():
            item = self.prompt_form_layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()

    def _add_prompt_sidebar_fallback(self, message: str) -> None:
        fallback_label = QLabel(message)
        fallback_label.setWordWrap(True)
        self.prompt_form_layout.addWidget(fallback_label)
        self.prompt_form_layout.addStretch(1)

    def _prompt_variable_config(self) -> dict[str, dict[str, object]]:
        return {
            "ten_cong_ty": {
                "label": "Tên công ty / Bối cảnh",
                "type": "combo",
                "options": [],
            },
            "vai_tro": {
                "label": "Vai trò của bạn",
                "type": "combo",
                "options": [],
            },
            "loai_nhiem_vu": {
                "label": "Loại nhiệm vụ",
                "type": "combo",
                "options": [
                    "Tạo nội dung",
                    "Phân tích dữ liệu",
                    "Lập kế hoạch/Chiến lược",
                    "Tối ưu hóa/Cải thiện",
                    "Dịch thuật/Chuyển đổi",
                    "Tóm tắt thông tin",
                ],
            },
            "giong_van": {
                "label": "Giọng văn",
                "type": "combo",
                "options": ["Chuyên nghiệp", "Thân thiện",  "Phân tích", "Lịch sự"],
            },
            "chuyen_mon": {
                "label": "Mức độ chuyên môn",
                "type": "combo",
                "options": ["Cơ bản", "Trung cấp", "Chuyên gia"],
            },
            "dinh_dang": {
                "label": "Định dạng kết quả",
                "type": "combo",
                "options": ["Markdown", "Email", "Báo cáo", "Bảng biểu", "Danh sách"],
            },
            "trinh_bay": {
                "label": "Yêu cầu trình bày",
                "type": "combo",
                "options": [
                    "Có tiêu đề rõ ràng, dùng bullet point",
                    "Trình bày dạng bảng so sánh",
                    "Phân tích từng bước chi tiết",
                ],
            },
            "gioi_han": {
                "label": "Giới hạn độ dài",
                "type": "combo",
                "options": ["Không giới hạn", "Dưới 500 từ", "Tối đa 3 mục chính", "Khoảng 1 trang A4"],
            },
        }

    def _add_prompt_field(self, placeholder: str, config: dict[str, object]) -> None:
        display_label = str(config.get("label") or placeholder)
        input_type = str(config.get("type") or "combo")

        label = QLabel(f"{display_label}:")
        label.setWordWrap(True)
        self.prompt_form_layout.addWidget(label)

        input_field = self._build_prompt_field_widget(placeholder, display_label, input_type, config)
        self.prompt_form_layout.addWidget(input_field)
        self.prompt_form_layout.addSpacing(6)
        self.prompt_field_inputs[placeholder] = input_field
        self._connect_prompt_field_autosave(input_field)

    def _build_prompt_field_widget(
        self,
        placeholder: str,
        display_label: str,
        input_type: str,
        config: dict[str, object],
    ) -> QWidget:
        if input_type == "text":
            input_field = QTextEdit()
            input_field.setObjectName("promptFieldInput")
            input_field.setPlaceholderText(str(config.get("placeholder") or ""))
            input_field.setFixedHeight(80)
            input_field.installEventFilter(self.wheel_event_filter)
            return input_field

        merged_options = self._merged_prompt_options(placeholder, config)
        input_field = self._create_prompt_input_widget(display_label, merged_options)
        input_field.installEventFilter(self.wheel_event_filter)
        return input_field

    def _merged_prompt_options(self, placeholder: str, config: dict[str, object]) -> list[str]:
        dynamic_options = self.prompt_options.get(placeholder, [])
        raw_default_options = config.get("options")
        if not isinstance(raw_default_options, list):
            raw_default_options = []

        default_options = [str(item) for item in raw_default_options if isinstance(item, str)]
        return list(dict.fromkeys(dynamic_options + default_options))

    def _add_company_context_checkbox_if_available(self) -> None:
        if not self.company_context_by_name:
            return

        self.company_context_checkbox = QCheckBox("Đính kèm thông tin công ty tham chiếu")
        self.company_context_checkbox.setChecked(False)
        self.company_context_checkbox.toggled.connect(self._on_right_panel_setting_changed)
        self.prompt_form_layout.addSpacing(4)
        self.prompt_form_layout.addWidget(self.company_context_checkbox)

    def _connect_prompt_field_autosave(self, input_field: QWidget) -> None:
        if isinstance(input_field, QLineEdit):
            input_field.textChanged.connect(self._on_right_panel_setting_changed)
            return

        if isinstance(input_field, QComboBox):
            input_field.currentTextChanged.connect(self._on_right_panel_setting_changed)
            return

        if isinstance(input_field, QTextEdit):
            input_field.textChanged.connect(self._on_right_panel_setting_changed)

    def _restore_prompt_sidebar_settings(self) -> None:
        self._restoring_right_panel_settings = True

        for placeholder, input_field in self.prompt_field_inputs.items():
            key = f"prompt_sidebar/field/{placeholder}"
            raw_value = self.settings.value(key)
            if raw_value is None:
                continue

            value = str(raw_value)
            if isinstance(input_field, QLineEdit):
                input_field.setText(value)
                continue

            if isinstance(input_field, QComboBox):
                input_field.setCurrentText(value)
                continue

            if isinstance(input_field, QTextEdit):
                input_field.setPlainText(value)

        if self.company_context_checkbox is not None:
            checkbox_raw = self.settings.value("prompt_sidebar/company_context_enabled", False)
            checkbox_value = self._coerce_setting_bool(checkbox_raw)
            self.company_context_checkbox.setChecked(checkbox_value)

        self._restoring_right_panel_settings = False

    def _coerce_setting_bool(self, value: object) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.strip().lower() in {"1", "true", "yes", "on"}
        if isinstance(value, (int, float)):
            return bool(value)
        return False

    def _on_right_panel_setting_changed(self, *_args) -> None:
        if self._restoring_right_panel_settings:
            return


        if self.search_grounding_checkbox is not None:
            self.settings.setValue(
                "chat/search_grounding_enabled",
                bool(self.search_grounding_checkbox.isChecked()),
            )

        if self.auto_open_export_checkbox is not None:
            self.settings.setValue(
                "export/auto_open_exported_files",
                bool(self.auto_open_export_checkbox.isChecked()),
            )

        if self.model_selector is not None:
            selected_model = self.model_selector.currentText().strip()
            if selected_model in self._MODEL_OPTIONS:
                self.fixed_model = selected_model
                self.settings.setValue("chat/model_name", selected_model)

        for placeholder, input_field in self.prompt_field_inputs.items():
            value = self._read_prompt_field_value(input_field)
            self.settings.setValue(f"prompt_sidebar/field/{placeholder}", value)

        if self.company_context_checkbox is not None:
            self.settings.setValue(
                "prompt_sidebar/company_context_enabled",
                bool(self.company_context_checkbox.isChecked()),
            )

    def _create_prompt_input_widget(self, display_label: str, options: list[str]) -> QWidget:
        combo = QComboBox()
        combo.setEditable(True)
        combo.setObjectName("promptFieldInput")
        combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        combo.setMaxVisibleItems(8)
        combo.addItems(options)
        
        # Cấu hình View để chặn cuộn ngang tuyệt đối
        view = combo.view()
        view.setTextElideMode(Qt.TextElideMode.ElideRight) # Tự động rút gọn text bằng dấu ...
        view.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff) # Tắt vĩnh viễn thanh cuộn ngang
        
        # Đồng bộ chiều rộng bảng chọn với chiều rộng của ô nhập (size sidebar)
        # Giúp bảng chọn không bị phình to hơn ô nhập.
        combo.setSizeAdjustPolicy(QComboBox.SizeAdjustPolicy.AdjustToMinimumContentsLengthWithIcon)
        
        # Style ép chiều rộng và ngăn tràn chữ
        combo.setStyleSheet("""
            QComboBox {
                max-width: 370px;
            }
            QComboBox QAbstractItemView {
                background-color: #ffffff;
                border: 1px solid #d1d5db;
                selection-background-color: #e5edff;
                selection-color: #1e3a8a;
                outline: 0px;
            }
            QComboBox QAbstractItemView::item {
                padding: 8px 10px;
                border-bottom: 1px solid #f3f4f6;
            }
        """)
        
        # Mặc định chọn cái đầu tiên nếu có
        if options:
            combo.setCurrentIndex(0)
        else:
            combo.setCurrentIndex(-1)
            combo.setEditText("")

        line_edit = combo.lineEdit()
        if line_edit is not None:
            line_edit.setPlaceholderText(f"Chọn hoặc nhập {display_label}...")
            line_edit.installEventFilter(self.wheel_event_filter)
        
        completer = combo.completer()
        if completer is not None:
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        return combo

    def _load_prompt_template_text(self) -> str:
        instructions_dir = get_instructions_dir()
        template_file = self._find_instruction_file_by_prefix(
            instructions_dir,
            self._TEMPLATE_INSTRUCTION_PREFIXES,
            recursive=True,
        )
        if template_file is not None:
            content = self._read_instruction_markdown_text(template_file)
            if content:
                return content

        for file_path in self._collect_instruction_markdown_files(instructions_dir, recursive=True):
            content = self._read_instruction_markdown_text(file_path)
            if content and "{{" in content and "}}" in content:
                return content

        return ""

    def _load_default_instruction_profile_text(self) -> str:
        instructions_dir = get_instructions_dir()
        base_text = self._load_instruction_file_text(
            instructions_dir,
            self._BASE_INSTRUCTION_PREFIXES,
        )
        overlay_text = self._load_instruction_file_text(
            instructions_dir,
            self._DEFAULT_OVERLAY_PREFIXES,
        )

        parts = [part for part in (base_text, overlay_text) if part]
        return "\n\n".join(parts).strip()

    def _load_instruction_file_text(self, directory: Path, candidate_prefixes: tuple[str, ...]) -> str:
        file_path = self._find_instruction_file_by_prefix(directory, candidate_prefixes, recursive=True)
        if file_path is None:
            return ""
        return self._read_instruction_markdown_text(file_path)

    def _collect_instruction_markdown_files(self, directory: Path, recursive: bool = True) -> list[Path]:
        if not directory.exists() or not directory.is_dir():
            return []

        if recursive:
            files = [path for path in directory.rglob("*.md") if path.is_file()]
        else:
            files = [path for path in directory.glob("*.md") if path.is_file()]

        return sorted(files, key=lambda path: str(path.relative_to(directory)).lower())

    def _find_instruction_file_by_prefix(
        self,
        directory: Path,
        prefixes: tuple[str, ...],
        recursive: bool = True,
    ) -> Path | None:
        normalized_prefixes = tuple(prefix.lower() for prefix in prefixes)

        for file_path in self._collect_instruction_markdown_files(directory, recursive=recursive):
            stem_lower = file_path.stem.lower()
            if any(stem_lower.startswith(prefix) for prefix in normalized_prefixes):
                return file_path

        return None

    def _read_instruction_markdown_text(self, file_path: Path) -> str:
        try:
            return file_path.read_text(encoding="utf-8").strip()
        except Exception:
            return ""

    def _extract_unique_placeholders(self, template_text: str) -> list[str]:
        placeholders = re.findall(r"{{(.*?)}}", template_text)
        seen: set[str] = set()
        result: list[str] = []
        for item in placeholders:
            item = item.strip()
            if not item or item in seen:
                continue
            seen.add(item)
            result.append(item)
        return result

    def _build_prompt_instructions(self) -> str | None:
        if not self.prompt_template_text:
            fallback = self.default_instructions_text.strip()
            return fallback or None

        result = self.prompt_template_text
        for placeholder in self._HIDDEN_PROMPT_KEYS:
            result = result.replace(f"{{{{{placeholder}}}}}", "")

        has_filled_value = False

        for placeholder, input_field in self.prompt_field_inputs.items():
            value = self._read_prompt_field_value(input_field)
            if value:
                has_filled_value = True

                if (
                    placeholder == self._PROMPT_COMPANY_KEY
                    and self.company_context_checkbox is not None
                    and self.company_context_checkbox.isChecked()
                ):
                    company_context = self._resolve_company_context(value)
                    if company_context:
                        value = f"{value}\n\nThông tin công ty tham chiếu:\n{company_context}"

                if placeholder == self._PROMPT_WORK_KEY:
                    # Nếu người dùng chọn từ danh sách, lấy nội dung file md đầy đủ
                    value = self.work_prompt_map.get(value, value)

                result = result.replace(f"{{{{{placeholder}}}}}", value)

        if has_filled_value:
            return result.strip()
        fallback = self.default_instructions_text.strip()
        return fallback or None

    def _read_prompt_field_value(self, input_field: QWidget) -> str:
        if isinstance(input_field, QLineEdit):
            return input_field.text().strip()
        if isinstance(input_field, QComboBox):
            return input_field.currentText().strip()
        if isinstance(input_field, QTextEdit):
            return input_field.toPlainText().strip()
        return ""

    def _resolve_company_context(self, company_name: str) -> str | None:
        if not company_name:
            return None

        direct_context = self.company_context_by_name.get(company_name)
        if direct_context:
            return direct_context

        return self.company_context_lookup.get(company_name.casefold())

    def _load_prompt_bundle_data(self) -> tuple[dict[str, list[str]], dict[str, str], dict[str, str], dict[str, str]]:
        sheets_dir = get_sheets_dir()
        
        # Load options from Cleaned_ .txt files if they exist, otherwise fallback to CSV
        company_txt = sheets_dir / "Cleaned_List_CongTy.txt"
        role_txt = sheets_dir / "Cleaned_List_VaiTro.txt"

        if company_txt.exists():
            company_options = [line.strip() for line in company_txt.read_text(encoding="utf-8").splitlines() if line.strip()]
        else:
            company_rows = self._read_csv_rows(sheets_dir / "List_CongTy.csv")
            company_options, _ = self._extract_company_options(company_rows)

        if role_txt.exists():
            role_options = [line.strip() for line in role_txt.read_text(encoding="utf-8").splitlines() if line.strip()]
        else:
            role_rows = self._read_csv_rows(sheets_dir / "List_VaiTro.csv")
            role_options = self._extract_role_options(role_rows)
        
        # We still need company_context for reference info even if using txt for dropdown names
        company_rows_all = self._read_csv_rows(sheets_dir / "List_CongTy.csv")
        _, company_context = self._extract_company_options(company_rows_all)

        # Load work options from markdown files in resources/instructions
        instructions_dir = get_instructions_dir()
        work_map = self._scan_work_instruction_files(instructions_dir)
        work_options = sorted(work_map.keys())

        company_lookup = {name.casefold(): info for name, info in company_context.items()}
        options = {
            self._PROMPT_COMPANY_KEY: company_options,
            self._PROMPT_ROLE_KEY: role_options,
            self._PROMPT_WORK_KEY: work_options,
        }
        return options, company_context, company_lookup, work_map

    def _scan_work_instruction_files(self, base_dir: Path) -> dict[str, str]:
        work_map = {}
        normalized_prefixes = {prefix.rstrip("_-") for prefix in self._WORK_INSTRUCTION_PREFIXES}
        prefix_pattern = re.compile(r"^(\d+)[_\-].+")

        for file_path in self._collect_instruction_markdown_files(base_dir, recursive=True):
            match = prefix_pattern.match(file_path.stem)
            if not match:
                continue

            group_prefix = match.group(1)
            if group_prefix not in normalized_prefixes:
                continue

            display_name = re.sub(r"[_\-]+", " ", file_path.stem).strip()
            display_key = f"[{group_prefix}] {display_name}"
            if display_key in work_map:
                relative_name = str(file_path.relative_to(base_dir)).replace("\\", "/")
                display_key = f"[{group_prefix}] {relative_name}"

            content = self._read_instruction_markdown_text(file_path)
            if content:
                work_map[display_key] = content

        return work_map


    def _read_csv_rows(self, csv_path: Path) -> list[dict[str, str]]:
        if not csv_path.exists() or not csv_path.is_file():
            return []

        for encoding in ("utf-8-sig", "utf-8"):
            try:
                with csv_path.open("r", encoding=encoding, newline="") as file:
                    reader = csv.DictReader(file)
                    rows: list[dict[str, str]] = []
                    for row in reader:
                        if not isinstance(row, dict):
                            continue

                        cleaned_row: dict[str, str] = {}
                        for key, value in row.items():
                            if key is None:
                                continue
                            normalized_key = str(key).strip()
                            cleaned_row[normalized_key] = value if isinstance(value, str) else ""
                        rows.append(cleaned_row)
                    return rows
            except Exception:
                continue

        return []

    def _extract_company_options(self, rows: object) -> tuple[list[str], dict[str, str]]:
        if not isinstance(rows, list):
            return [], {}

        options: list[str] = []
        contexts: dict[str, str] = {}
        seen: set[str] = set()

        for row in rows:
            if not isinstance(row, dict):
                continue

            company_name = self._clean_text(
                row.get("TÊN CÔNG TY") or row.get("col_3") or row.get("Name")
            )
            if not self._is_valid_option(company_name, max_len=140):
                continue

            if company_name not in seen:
                seen.add(company_name)
                options.append(company_name)

            if company_name in contexts:
                continue

            context_text = self._clean_company_context(
                row.get("THÔNG TIN CÔNG TY") or row.get("col_6") or row.get("col_4")
            )
            if context_text:
                contexts[company_name] = context_text

        return options, contexts

    def _extract_role_options(self, rows: object) -> list[str]:
        if not isinstance(rows, list):
            return []

        options: list[str] = []
        seen: set[str] = set()

        for row in rows:
            if not isinstance(row, dict):
                continue

            raw_text = row.get("VAI TRÒ, VỊ TRÍ CÔNG VIỆC (Ô B1 input)") or row.get("col_3")
            role_text = self._clean_text(raw_text)

            if not self._is_valid_option(role_text, max_len=120):
                continue

            if role_text.isupper() and len(role_text) > 20:
                continue

            if role_text not in seen:
                seen.add(role_text)
                options.append(role_text)

        return options

    def _clean_text(self, value: object) -> str:
        if not isinstance(value, str):
            return ""

        text = value.replace("\xa0", " ").strip()
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _clean_company_context(self, value: object) -> str:
        if not isinstance(value, str):
            return ""

        text = value.replace("```", "").strip()
        if text.startswith("---"):
            text = text[3:].lstrip()

        marker = "Bạn có thể khai thác thêm thông tin về"
        if marker in text and "\n" in text:
            _, text = text.split("\n", 1)

        text = text.strip()
        if len(text) < 40:
            return ""
        if len(text) > 3200:
            text = f"{text[:3200].rstrip()}..."

        return text

    def _is_valid_option(self, text: str, max_len: int = 180) -> bool:
        if not text:
            return False
        if len(text) < 2 or len(text) > max_len:
            return False
        if text.startswith("```") or text.startswith("---"):
            return False
        if "Bạn có thể khai thác thêm thông tin về" in text:
            return False

        char_set = set(text)
        if char_set and char_set.issubset({"_", "-", ".", ",", ":", ";", " ", "|"}):
            return False

        return True

    def _load_conversations(self) -> None:
        selected_id = self.state.current_conversation_id
        try:
            conversations = self.client.list_conversations()
        except ApiError as exc:
            self._show_error(str(exc))
            return

        self.conversation_list.blockSignals(True)
        self.conversation_list.clear()
        for conversation in conversations:
            label = self._conversation_label(conversation)
            item = QListWidgetItem(label)
            item.setData(Qt.ItemDataRole.UserRole, conversation.id)
            item.setData(Qt.ItemDataRole.UserRole + 1, conversation.title or "")
            item.setToolTip(conversation.title or label)
            self.conversation_list.addItem(item)
            if selected_id and conversation.id == selected_id:
                item.setSelected(True)
        self.conversation_list.blockSignals(False)

    def _on_conversation_selected(self) -> None:
        current_item = self.conversation_list.currentItem()
        if current_item is None:
            return
        conversation_id = current_item.data(Qt.ItemDataRole.UserRole)
        if not isinstance(conversation_id, str):
            return
        self._load_history(conversation_id)

    def _on_conversation_context_menu(self, position) -> None:
        item = self.conversation_list.itemAt(position)
        if item is None:
            return

        self.conversation_list.setCurrentItem(item)

        conversation_id = item.data(Qt.ItemDataRole.UserRole)
        if not isinstance(conversation_id, str):
            return

        menu = QMenu(self)
        rename_action = menu.addAction("Đổi tên cuộc trò chuyện")
        delete_action = menu.addAction("Xóa cuộc trò chuyện")
        chosen_action = menu.exec(self.conversation_list.viewport().mapToGlobal(position))
        if chosen_action == rename_action:
            self._rename_conversation(conversation_id)
        if chosen_action == delete_action:
            self._delete_conversation(conversation_id)

    def _rename_conversation(self, conversation_id: str) -> None:
        current_item = self.conversation_list.currentItem()
        current_title = ""
        if current_item is not None:
            maybe_title = current_item.data(Qt.ItemDataRole.UserRole + 1)
            if isinstance(maybe_title, str):
                current_title = maybe_title

        new_title, ok = QInputDialog.getText(
            self,
            "Đổi tên cuộc trò chuyện",
            "Tên mới:",
            text=current_title,
        )
        if not ok:
            return

        cleaned_title = new_title.strip()
        if not cleaned_title:
            self._show_error("Tên cuộc trò chuyện không được để trống.")
            return

        try:
            self.client.rename_conversation(conversation_id, cleaned_title)
        except ApiError as exc:
            self._show_error(str(exc))
            return

        self._load_conversations()
        self.statusBar().showMessage("Đã đổi tên cuộc trò chuyện", 3000)

    def _delete_conversation(self, conversation_id: str) -> None:
        answer = QMessageBox.question(
            self,
            "Xóa cuộc trò chuyện",
            "Bạn có chắc muốn xóa cuộc trò chuyện này và toàn bộ tin nhắn?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if answer != QMessageBox.StandardButton.Yes:
            return

        try:
            self.client.delete_conversation(conversation_id)
        except ApiError as exc:
            self._show_error(str(exc))
            return

        if self.state.current_conversation_id == conversation_id:
            self.state.reset_chat()
            self._render_messages()
            self._update_attachment_label()

        self._load_conversations()
        self.statusBar().showMessage("Đã xóa cuộc trò chuyện", 3000)

    def _load_history(self, conversation_id: str) -> None:
        previous_conversation_id = self.state.current_conversation_id

        try:
            history = self.client.get_history(conversation_id)
        except ApiError as exc:
            self._show_error(str(exc))
            return

        if previous_conversation_id and previous_conversation_id != conversation_id:
            self.state.attached_paths.clear()
            self._update_attachment_label()

        self.state.current_conversation_id = conversation_id
        messages = [
            ChatMessage(
                role=msg.role,
                text=self._extract_text(msg),
                created_at=msg.created_at,
            )
            for msg in history
        ]
        self.state.set_messages(messages)
        self._render_messages()

    def _send_message(self) -> None:
        if self.stream_worker and self.stream_worker.isRunning():
            return

        prompt = self.input_box.toPlainText().strip()
        if not prompt:
            return

        conversation_id = self.state.current_conversation_id or str(uuid.uuid4())
        self.state.current_conversation_id = conversation_id

        request = ChatRequest(
            conversation_id=conversation_id,
            instructions=self._resolve_request_instructions(),
            input=prompt,
            model=self._selected_model(),
            file_paths=list(self.state.attached_paths),
            search_grounding=self.search_grounding_checkbox.isChecked() if self.search_grounding_checkbox is not None else True,
        )

        attachment_names = [Path(path).name for path in self.state.attached_paths]

        self.state.add_message(role="user", text=prompt, attachment_names=attachment_names)
        self.state.add_message(role="assistant", text="")
        self._render_messages()

        self.input_box.clear()
        self._set_busy_state(True)
        self._set_response_status("Trạng thái phản hồi: Đang tạo phản hồi...", "processing")
        self.statusBar().showMessage("Đang tạo phản hồi...")

        self.stream_worker = ChatStreamWorker(self.client, request)
        self.stream_worker.success.connect(self._on_stream_success)
        self.stream_worker.failed.connect(self._on_stream_failed)
        self.stream_worker.finished.connect(self._on_stream_finished)
        self.stream_worker.start()

    def _on_stream_success(self, result: StreamResult) -> None:
        if self.state.messages and self.state.messages[-1].role == "assistant":
            self.state.messages[-1].text = result.text
        else:
            self.state.add_message(role="assistant", text=result.text)

        self.state.current_conversation_id = result.conversation_id
        self._render_messages()

        self._load_conversations()
        finished_at = datetime.now().strftime("%H:%M:%S")
        self._set_response_status(f"Trạng thái phản hồi: Đã hoàn tất lúc {finished_at}", "done")
        if self.state.attached_paths:
            self.statusBar().showMessage(
                f"Hoàn tất ({result.status}) • giữ lại {len(self.state.attached_paths)} tệp đính kèm",
                4000,
            )
        else:
            self.statusBar().showMessage(f"Hoàn tất ({result.status})", 4000)

    def _on_stream_failed(self, error_message: str) -> None:
        if self.state.messages and self.state.messages[-1].role == "assistant" and not self.state.messages[-1].text:
            self.state.messages.pop()
        self._render_messages()
        self._set_response_status("Trạng thái phản hồi: Có lỗi khi tạo phản hồi", "error")
        self._show_error(error_message)

    def _on_stream_finished(self) -> None:
        self._set_busy_state(False)

    def _set_response_status(self, text: str, state: str) -> None:
        # Update internal status and re-render messages so the assistant bubble shows it
        self._response_status_text = text
        self._response_status_state = state
        self._render_messages()

    def _adjust_input_box_height(self) -> None:
        document_height = int(self.input_box.document().size().height())
        target_height = max(self._INPUT_MIN_HEIGHT, min(document_height + 14, self._INPUT_MAX_HEIGHT))
        self.input_box.setFixedHeight(target_height)

        needs_scroll = document_height + 14 > self._INPUT_MAX_HEIGHT
        self.input_box.setVerticalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded
            if needs_scroll
            else Qt.ScrollBarPolicy.ScrollBarAlwaysOff
        )

    def _configure_icon_button(
        self,
        button: QPushButton,
        theme_icon_names: tuple[str, ...],
        fallback_symbol: str,
        tooltip: str,
        icon_size: int = 18,
    ) -> None:
        resolved_icon = QIcon()
        for icon_name in theme_icon_names:
            candidate_icon = QIcon.fromTheme(icon_name)
            if not candidate_icon.isNull():
                resolved_icon = candidate_icon
                break

        if resolved_icon.isNull():
            button.setText(fallback_symbol)
        else:
            source_pixmap = resolved_icon.pixmap(icon_size, icon_size)
            if not source_pixmap.isNull():
                tinted_pixmap = source_pixmap.copy()
                painter = QPainter(tinted_pixmap)
                painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_SourceIn)
                painter.fillRect(tinted_pixmap.rect(), QColor("#ffffff"))
                painter.end()
                resolved_icon = QIcon(tinted_pixmap)
            button.setIcon(resolved_icon)
            button.setText("")

        button.setIconSize(QSize(icon_size, icon_size))
        button.setToolTip(tooltip)

    def _configure_resource_svg_icon_button(
        self,
        button: QPushButton,
        svg_name: str,
        theme_icon_names: tuple[str, ...],
        fallback_symbol: str,
        tooltip: str,
        icon_size: int = 18,
    ) -> None:
        icon_dir = get_icons_dir()
        svg_path = icon_dir / svg_name

        resolved_icon = QIcon()
        if svg_path.exists():
            resolved_icon = QIcon(str(svg_path))

        if resolved_icon.isNull():
            self._configure_icon_button(
                button,
                theme_icon_names,
                fallback_symbol,
                tooltip,
                icon_size=icon_size,
            )
            return

        source_pixmap = resolved_icon.pixmap(icon_size, icon_size)
        if not source_pixmap.isNull():
            tinted_pixmap = source_pixmap.copy()
            painter = QPainter(tinted_pixmap)
            painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_SourceIn)
            painter.fillRect(tinted_pixmap.rect(), QColor("#ffffff"))
            painter.end()
            resolved_icon = QIcon(tinted_pixmap)

        button.setIcon(resolved_icon)
        button.setText("")
        button.setIconSize(QSize(icon_size, icon_size))
        button.setToolTip(tooltip)

    def _set_busy_state(self, is_busy: bool) -> None:
        self.input_box.setEnabled(not is_busy)
        self.send_button.setEnabled(not is_busy)
        self.add_file_button.setEnabled(not is_busy)
        self.add_image_button.setEnabled(not is_busy)
        if self.model_selector is not None:
            self.model_selector.setEnabled(not is_busy)

        if is_busy:
            self.clear_file_button.setEnabled(False)

            self.input_box.setPlaceholderText("Đang tạo phản hồi...")
            self._start_response_spinner()
            return

        self._update_attachment_label()
        self.input_box.setPlaceholderText("Nhập tin nhắn...")
        self._stop_response_spinner()

    def _selected_model(self) -> str:
        if self.model_selector is not None:
            selected_model = self.model_selector.currentText().strip()
            if selected_model in self._MODEL_OPTIONS:
                self.fixed_model = selected_model
                return selected_model

        return self.fixed_model

    def _start_response_spinner(self) -> None:
        # Start the spinner timer and reset index; rendering of the spinner
        # happens inside the assistant bubble via _render_messages.
        self._spinner_index = 0
        if not self._spinner_timer.isActive():
            self._spinner_timer.start()
        self._render_messages()

    def _stop_response_spinner(self) -> None:
        # Stop spinner timer and refresh assistant bubble rendering
        self._spinner_timer.stop()
        self._render_messages()

    def _advance_response_spinner(self) -> None:
        if not self._spinner_timer.isActive():
            return

        self._spinner_index = (self._spinner_index + 1) % len(self._spinner_frames)
        # refresh chat view so assistant bubble shows updated spinner frame
        self._render_messages()

    def _attach_files(self) -> None:
        default_dir = Path.home() / "Desktop"
        if not default_dir.exists():
            default_dir = Path.home()

        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Chọn tệp",
            str(default_dir),
        )
        if not file_paths:
            return
        known = set(self.state.attached_paths)
        for path in file_paths:
            if path not in known:
                self.state.attached_paths.append(path)
                known.add(path)
        self._update_attachment_label()

    def _attach_images(self) -> None:
        default_dir = Path.home() / "Desktop"
        if not default_dir.exists():
            default_dir = Path.home()

        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Chọn ảnh",
            str(default_dir),
            "Image Files (*.png *.jpg *.jpeg *.webp *.bmp *.gif *.tif *.tiff)",
        )
        if not file_paths:
            return

        known = set(self.state.attached_paths)
        for path in file_paths:
            if path not in known:
                self.state.attached_paths.append(path)
                known.add(path)
        self._update_attachment_label()

    def _clear_attachments(self) -> None:
        self.state.attached_paths.clear()
        self._update_attachment_label()

    def _update_attachment_label(self) -> None:
        if not self.state.attached_paths:
            self.attachment_list.setVisible(False)
            self.clear_file_button.setEnabled(False)
            return

        self.attachment_list.setVisible(True)
        self.clear_file_button.setEnabled(True)

        html_blocks = [
            "<html><body style='margin:0; padding:8px; font-family:Segoe UI, sans-serif; background-color:#f8fafc;'>",
            "<table cellspacing='10' cellpadding='0'><tr>"
        ]

        for file_path in self.state.attached_paths:
            path_obj = Path(file_path)
            file_name = path_obj.name
            ext = path_obj.suffix[1:].upper() if path_obj.suffix else "FILE"
            
            # Chọn màu theo loại file
            color = "#3b82f6" # Blue
            if ext in ["PDF"]: color = "#ef4444" # Red
            elif ext in ["DOC", "DOCX"]: color = "#2563eb" # Dark Blue
            elif ext in ["XLS", "XLSX", "CSV"]: color = "#10b981" # Green
            elif ext in ["PNG", "JPG", "JPEG", "WEBP"]: color = "#8b5cf6" # Purple

            # Nội dung icon/thumbnail
            if self._is_image_attachment_path(file_path):
                file_url = QUrl.fromLocalFile(file_path).toString()
                content = f"<img src='{file_url}' width='50' height='50' style='border-radius:4px;'>"
            else:
                content = (
                    f"<div style='width:50px; height:50px; background:#f1f5f9; color:{color}; "
                    "font-weight:bold; font-size:11px; border-radius:4px; border:1px solid #e2e8f0; "
                    "text-align:center; padding-top:16px;'>"
                    f"{ext}</div>"
                )

            html_blocks.append(
                f"<td valign='top'>"
                f"<div style='width:80px; background:white; border:1px solid #cbd5e1; border-radius:8px;'>"
                f"<div style='height:4px; background-color:{color}; border-top-left-radius:8px; border-top-right-radius:8px;'></div>"
                f"<div style='padding:8px; text-align:center;'>"
                f"{content}"
                f"<div style='margin-top:6px; font-size:10px; color:#1e293b; height:24px; overflow:hidden;'>{html.escape(file_name)}</div>"
                f"</div></div>"
                f"</td>"
            )

        html_blocks.append("</tr></table></body></html>")
        self.attachment_list.setHtml("".join(html_blocks))
        self.attachment_list.setToolTip("\n".join(self.state.attached_paths))

    def _is_image_attachment_path(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self._IMAGE_EXTENSIONS

    def _resolve_request_instructions(self) -> str | None:
        has_image_file = any(
            self._is_image_attachment_path(file_path)
            for file_path in self.state.attached_paths
        )
        if has_image_file:
            return None

        has_non_image_file = any(
            not self._is_image_attachment_path(file_path)
            for file_path in self.state.attached_paths
        )
        if not has_non_image_file:
            return None

        return self._build_prompt_instructions()

    def _attachment_chip_colors(self, file_name: str) -> tuple[str, str, str]:
        extension = Path(file_name).suffix.lower()

        if extension in self._IMAGE_EXTENSIONS:
            return "#bfdbfe", "#eff6ff", "#1e3a8a"
        if extension == ".pdf":
            return "#fecaca", "#fef2f2", "#991b1b"
        if extension in {".doc", ".docx"}:
            return "#c7d2fe", "#eef2ff", "#3730a3"
        if extension in {".xls", ".xlsx", ".csv"}:
            return "#bbf7d0", "#f0fdf4", "#166534"
        if extension in {".txt", ".md", ".json", ".xml", ".yaml", ".yml"}:
            return "#d1d5db", "#f9fafb", "#374151"
        return "#d1d5db", "#f3f4f6", "#374151"

    def _build_attachment_chip_html(self, file_name: str) -> str:
        border_color, background_color, text_color = self._attachment_chip_colors(file_name)
        return (
            "<span style='display:inline-block; margin:0 6px 6px 0; padding:3px 8px; "
            "border-radius:999px; "
            f"border:1px solid {border_color}; background:{background_color}; color:{text_color}; "
            "font-size:12px;'>"
            f"{html.escape(file_name)}"
            "</span>"
        )

    def _render_messages(self) -> None:
        blocks: list[str] = [
            (
                "<html><head>"
                "<style>"
                "body { margin: 10px; font-family: 'Segoe UI', sans-serif; font-size: 13px; color: #111827; background-color: #ffffff; }"
                ".message-container { margin-bottom: 15px; width: 100%; }"
                ".user-bubble { background-color: #007bff; color: white; border-radius: 15px 15px 0 15px; padding: 10px 15px; margin-left: 50px; border: 1px solid #0069d9; }"
                ".assistant-bubble { background-color: #f1f3f4; color: #202124; border-radius: 15px 15px 15px 0; padding: 12px 18px; margin-right: 50px; border: 1px solid #e0e0e0; }"
                ".meta { font-size: 11px; color: #70757a; margin-bottom: 4px; font-weight: 600; }"
                ".user-meta { text-align: right; color: #007bff; }"
                ".timestamp { font-size: 10px; color: rgba(0,0,0,0.4); margin-top: 5px; }"
                ".user-timestamp { color: rgba(255,255,255,0.7); text-align: right; }"
                "code { background-color: rgba(0,0,0,0.05); padding: 2px 4px; border-radius: 4px; font-family: 'Consolas', monospace; }"
                "pre { background-color: #2d2d2d; color: #f8f8f2; padding: 10px; border-radius: 8px; font-family: 'Consolas', monospace; }"
                "table { border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 12px; }"
                "th, td { border: 1px solid #dfe1e5; padding: 6px; text-align: left; }"
                "th { background-color: #f8f9fa; }"
                "a { color: #007bff; text-decoration: none; font-weight: bold; }"
                ".assistant-bubble a { color: #1a73e8; }"
                "</style>"
                "</head><body>"
            )
        ]

        latest_assistant_index = -1
        for idx in range(len(self.state.messages) - 1, -1, -1):
            if self.state.messages[idx].role.lower() == "assistant":
                latest_assistant_index = idx
                break

        for message_index, message in enumerate(self.state.messages):
            role = message.role.lower()
            is_user = (role == "user")
            title = "Bạn" if is_user else "Trợ lý"
            text = self._render_markdown_html(message.text)
            attachments_html = ""

            if is_user and message.attachment_names:
                filename_items = "".join(
                    self._build_attachment_chip_html(file_name)
                    for file_name in message.attachment_names
                )
                attachments_html = (
                    "<div style='margin-top:8px; padding:6px 8px; border-radius:8px; "
                    "background:rgba(255,255,255,0.2); border:1px solid rgba(255,255,255,0.3);'>"
                    "<div style='font-size:11px; font-weight:700; color:white; margin-bottom:4px;'>"
                    "Tệp đính kèm:</div>"
                    f"<div style='margin:0;'>{filename_items}</div>"
                    "</div>"
                )

            timestamp_str = ""
            if isinstance(message.created_at, datetime):
                timestamp_str = message.created_at.strftime("%H:%M")

            actions_html = ""
            assistant_status_html = ""
            if not is_user and message.text.strip():
                actions_html = (
                    "<div style='margin-top:10px;'>"
                    f"<a href='action://export-word/{message_index}'>📄 Xuất Word</a>"
                    "&nbsp;&nbsp;&nbsp;"
                    f"<a href='action://export-pdf/{message_index}'>📕 Xuất PDF</a>"
                    "</div>"
                )

            if not is_user and message_index == latest_assistant_index:
                # ... (spinner logic)
                state_label_map = {
                    "processing": "Đang phản hồi...",
                    "done": "Đã phản hồi",
                    "error": "Phản hồi lỗi",
                    "idle": "Sẵn sàng",
                }
                status_label = state_label_map.get(self._response_status_state, self._response_status_text)
                spinner_html = ""
                if self._response_status_state == "processing":
                    frame = self._spinner_frames[self._spinner_index] if self._spinner_timer.isActive() else self._spinner_frames[0]
                    spinner_html = f"<span style='color:#1a73e8; font-weight:bold;'>{html.escape(frame)} </span>"

                assistant_status_html = (
                    f"<div style='margin-bottom:8px; font-size:11px; font-weight:bold; color:#5f6368;'>"
                    f"{spinner_html}Trạng thái: {status_label}"
                    "</div>"
                )

            if is_user:
                bubble = (
                    "<div class='message-container'>"
                    f"<div class='meta user-meta'>{title}</div>"
                    f"<div class='user-bubble'>"
                    f"{text}"
                    f"{attachments_html}"
                    f"<div class='timestamp user-timestamp'>{timestamp_str}</div>"
                    "</div></div>"
                )
            else:
                bubble = (
                    "<div class='message-container'>"
                    f"<div class='meta'>{title}</div>"
                    f"<div class='assistant-bubble'>"
                    f"{assistant_status_html}"
                    f"{text}"
                    f"{actions_html}"
                    f"<div class='timestamp'>{timestamp_str}</div>"
                    "</div></div>"
                )
            blocks.append(bubble)

        blocks.append("</body></html>")
        self.chat_view.setHtml("".join(blocks))
        self.chat_view.verticalScrollBar().setValue(self.chat_view.verticalScrollBar().maximum())

    def _extract_text(self, message: BaseMessage) -> str:
        content_text = message.content.get("text", "")
        if isinstance(content_text, str):
            return content_text
        return str(content_text)

    def _on_chat_link_clicked(self, url: QUrl) -> None:
        if url.scheme() != "action":
            return

        action = url.host().strip().lower()
        raw_index = url.path().lstrip("/")
        try:
            message_index = int(raw_index)
        except ValueError:
            self._show_error("Liên kết thao tác không hợp lệ.")
            return

        if message_index < 0 or message_index >= len(self.state.messages):
            self._show_error("Không tìm thấy phản hồi cần xuất.")
            return

        message = self.state.messages[message_index]
        if message.role.lower() != "assistant" or not message.text.strip():
            self._show_error("Chỉ có thể xuất phản hồi của Trợ lý.")
            return

        if action == "export-word":
            self._export_assistant_message_to_word(message)
            return

        if action == "export-pdf":
            self._export_assistant_message_to_pdf(message)
            return

        self._show_error("Thao tác chưa được hỗ trợ.")

    def _render_markdown_html(self, text: str) -> str:
        normalized = text.strip()
        if not normalized:
            return ""

        document = QTextDocument()
        document.setMarkdown(normalized)
        rendered = document.toHtml()

        body_start = rendered.find("<body")
        if body_start < 0:
            return html.escape(text).replace("\n", "<br/>")

        body_open_end = rendered.find(">", body_start)
        body_end = rendered.rfind("</body>")
        if body_open_end < 0 or body_end <= body_open_end:
            return html.escape(text).replace("\n", "<br/>")

        body_html = rendered[body_open_end + 1:body_end].strip()
        if not body_html:
            return html.escape(text).replace("\n", "<br/>")

        return body_html

    def _export_conversation_to_word(self) -> None:
        assistant_messages = [
            message
            for message in self.state.messages
            if message.role.lower() == "assistant" and message.text.strip()
        ]
        if not assistant_messages:
            QMessageBox.information(
                self,
                "Xuất Word",
                "Chưa có phản hồi Trợ lý để xuất.",
            )
            return

        latest_assistant_message = assistant_messages[-1]

        self._export_assistant_message_to_word(latest_assistant_message)

    def _export_assistant_message_to_word(self, assistant_message: ChatMessage) -> None:
        if assistant_message.role.lower() != "assistant" or not assistant_message.text.strip():
            QMessageBox.information(
                self,
                "Xuất Word",
                "Phản hồi Trợ lý không có nội dung để xuất.",
            )
            return

        try:
            from docx import Document
        except Exception:
            self._show_error(
                "Thiếu thư viện xuất Word. Hãy cài `python-docx` rồi thử lại."
            )
            return

        title = self._current_conversation_title() or "Cuộc trò chuyện"
        default_stem = self._safe_filename(f"{title}_tro_ly_{datetime.now().strftime('%Y%m%d_%H%M')}")
        default_dir = Path.home() / "Desktop"
        if not default_dir.exists():
            default_dir = Path.home()
        default_path = default_dir / f"{default_stem}.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Lưu file Word",
            str(default_path),
            "Word Document (*.docx)",
        )
        if not file_path:
            return

        output_path = Path(file_path)
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")

        markdown_text = assistant_message.text.strip()
        if not markdown_text:
            QMessageBox.information(
                self,
                "Xuất Word",
                "Phản hồi Trợ lý không có nội dung để xuất.",
            )
            return

        document = Document()
        self._apply_word_document_style(document)
        self._append_markdown_to_word_document(document, markdown_text)

        try:
            document.save(str(output_path))
        except Exception as exc:
            self._show_error(f"Không thể lưu file Word: {exc}")
            return

        self.statusBar().showMessage(f"Đã xuất 1 phản hồi Trợ lý: {output_path.name}", 4000)
        if self.auto_open_export_checkbox is None or self.auto_open_export_checkbox.isChecked():
            self._open_exported_file(output_path)

    def _apply_word_document_style(self, document) -> None:
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt

        for section in document.sections:
            section.top_margin = Cm(2.3)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(11)
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = 1.25
        normal_paragraph.space_after = Pt(6)

        if "Heading 1" in document.styles:
            heading_1 = document.styles["Heading 1"]
            heading_1.font.name = "Times New Roman"
            heading_1.font.size = Pt(16)
            heading_1.font.bold = True

        if "Heading 2" in document.styles:
            heading_2 = document.styles["Heading 2"]
            heading_2.font.name = "Times New Roman"
            heading_2.font.size = Pt(14)
            heading_2.font.bold = True

    def _normalize_export_markdown_text(self, text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.replace("\u00a0", " ").replace("\u200b", "")

        normalized = re.sub(r"\t+", " ", normalized)
        normalized = re.sub(r"(?<!\d)([;:!?])(?!\s)(?=\S)", r"\1 ", normalized)
        normalized = re.sub(r"(?<!\d)([.,])(?!\d)(?!\s)(?=\S)", r"\1 ", normalized)
        normalized = re.sub(r"([\]})])([A-Za-zÀ-Ỵà-ỵ0-9])", r"\1 \2", normalized)
        normalized = re.sub(r"([A-Za-zÀ-Ỵà-ỵ0-9])([\[({])", r"\1 \2", normalized)
        normalized = re.sub(r"[ ]{2,}", " ", normalized)

        lines = [line.strip() for line in normalized.split("\n")]
        compact: list[str] = []
        previous_blank = False
        for line in lines:
            if not line:
                if not previous_blank:
                    compact.append("")
                previous_blank = True
                continue
            compact.append(line)
            previous_blank = False
        return "\n".join(compact).strip()

    def _split_long_text_for_word(self, text: str) -> list[str]:
        if len(text) < 450:
            return [text]

        parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÀ-Ỵ0-9])", text)
        cleaned = [part.strip() for part in parts if part.strip()]
        if len(cleaned) <= 1:
            return [text]
        return cleaned

    def _is_markdown_table_separator(self, line: str) -> bool:
        return bool(re.match(r"^\|?\s*:?[-]{3,}:?(\s*\|\s*:?[-]{3,}:?)*\s*\|?$", line))

    def _append_markdown_table_to_document(self, document, lines: list[str]) -> None:
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        if len(cleaned_lines) < 2:
            return

        header_line = cleaned_lines[0]
        separator_line = cleaned_lines[1]
        if not self._is_markdown_table_separator(separator_line):
            return

        data_lines = cleaned_lines[2:]
        header_cells = [cell.strip() for cell in header_line.strip("|").split("|")]
        if not header_cells:
            return

        table = document.add_table(rows=1, cols=len(header_cells))
        table.style = "Table Grid"

        for index, value in enumerate(header_cells):
            header_paragraph = table.rows[0].cells[index].paragraphs[0]
            self._add_markdown_runs_to_paragraph(header_paragraph, value)
            for run in header_paragraph.runs:
                run.bold = True

        for line in data_lines:
            row_values = [cell.strip() for cell in line.strip("|").split("|")]
            row = table.add_row()
            for index in range(len(header_cells)):
                paragraph = row.cells[index].paragraphs[0]
                value = row_values[index] if index < len(row_values) else ""
                self._add_markdown_runs_to_paragraph(paragraph, value)

        document.add_paragraph("")

    def _flush_word_paragraph_lines(self, document, paragraph_lines: list[str]) -> None:
        from docx.shared import Pt

        if not paragraph_lines:
            return

        merged_text = re.sub(r"\s+", " ", " ".join(paragraph_lines)).strip()
        paragraph_lines.clear()
        if not merged_text:
            return

        for chunk in self._split_long_text_for_word(merged_text):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            self._add_markdown_runs_to_paragraph(paragraph, chunk)

    def _flush_word_table_lines(self, document, table_lines: list[str]) -> None:
        if not table_lines:
            return

        self._append_markdown_table_to_document(document, list(table_lines))
        table_lines.clear()

    def _append_word_code_block(self, document, code_lines: list[str]) -> None:
        from docx.shared import Pt

        if not code_lines:
            return

        paragraph = document.add_paragraph("\n".join(code_lines))
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        code_lines.clear()

    def _append_word_heading_line(self, document, stripped: str) -> bool:
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if not heading_match:
            return False

        level = min(len(heading_match.group(1)), 4)
        heading_text = heading_match.group(2).strip()
        document.add_heading(heading_text, level=level)
        return True

    def _append_word_bullet_line(self, document, stripped: str) -> bool:
        from docx.shared import Pt

        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if not bullet_match:
            return False

        content = bullet_match.group(1).strip()
        try:
            paragraph = document.add_paragraph(style="List Bullet")
        except Exception:
            paragraph = document.add_paragraph("• ")
        self._add_markdown_runs_to_paragraph(paragraph, content)
        paragraph.paragraph_format.space_after = Pt(4)
        return True

    def _append_word_numbered_line(self, document, stripped: str) -> bool:
        from docx.shared import Pt

        numbered_match = re.match(r"^(\d+)[\.)]\s+(.+)$", stripped)
        if not numbered_match:
            return False

        content = numbered_match.group(2).strip()
        try:
            paragraph = document.add_paragraph(style="List Number")
        except Exception:
            paragraph = document.add_paragraph(f"{numbered_match.group(1)}. ")
        self._add_markdown_runs_to_paragraph(paragraph, content)
        paragraph.paragraph_format.space_after = Pt(4)
        return True

    def _append_word_quote_line(self, document, stripped: str) -> bool:
        from docx.shared import Pt

        quote_match = re.match(r"^>\s?(.*)$", stripped)
        if not quote_match:
            return False

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(14)
        paragraph.paragraph_format.space_after = Pt(6)
        self._add_markdown_runs_to_paragraph(paragraph, quote_match.group(1).strip())
        return True

    def _append_markdown_to_word_document(self, document, text: str) -> None:
        normalized_text = self._normalize_export_markdown_text(text)
        in_code_block = False
        code_lines: list[str] = []
        paragraph_lines: list[str] = []
        table_lines: list[str] = []

        for raw_line in normalized_text.splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()

            if stripped.startswith("```"):
                self._flush_word_paragraph_lines(document, paragraph_lines)
                self._flush_word_table_lines(document, table_lines)
                if in_code_block:
                    self._append_word_code_block(document, code_lines)
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if not stripped:
                self._flush_word_paragraph_lines(document, paragraph_lines)
                self._flush_word_table_lines(document, table_lines)
                document.add_paragraph("")
                continue

            if "|" in stripped:
                self._flush_word_paragraph_lines(document, paragraph_lines)
                table_lines.append(stripped)
                continue

            if table_lines:
                self._flush_word_table_lines(document, table_lines)

            if self._append_word_heading_line(document, stripped):
                continue

            if self._append_word_bullet_line(document, stripped):
                continue

            if self._append_word_numbered_line(document, stripped):
                continue

            if self._append_word_quote_line(document, stripped):
                continue

            paragraph_lines.append(stripped)

        self._flush_word_paragraph_lines(document, paragraph_lines)
        self._flush_word_table_lines(document, table_lines)

        if in_code_block and code_lines:
            self._append_word_code_block(document, code_lines)

    def _add_markdown_runs_to_paragraph(self, paragraph, text: str) -> None:
        normalized_text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
        normalized_text = normalized_text.replace("\\*", "*").replace("\\_", "_")
        token_pattern = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_)")

        for token in token_pattern.split(normalized_text):
            if not token:
                continue

            if token.startswith("**") and token.endswith("**") and len(token) >= 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
                continue

            if token.startswith("__") and token.endswith("__") and len(token) >= 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
                continue

            if token.startswith("`") and token.endswith("`") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
                continue

            if token.startswith("*") and token.endswith("*") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
                continue

            if token.startswith("_") and token.endswith("_") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
                continue

            paragraph.add_run(token)

    def _export_assistant_message_to_pdf(self, assistant_message: ChatMessage) -> None:
        if assistant_message.role.lower() != "assistant" or not assistant_message.text.strip():
            QMessageBox.information(
                self,
                "Xuất PDF",
                "Phản hồi Trợ lý không có nội dung để xuất.",
            )
            return

        title = self._current_conversation_title() or "Cuộc trò chuyện"
        default_stem = self._safe_filename(f"{title}_tro_ly_{datetime.now().strftime('%Y%m%d_%H%M')}")
        default_dir = Path.home() / "Desktop"
        if not default_dir.exists():
            default_dir = Path.home()
        default_path = default_dir / f"{default_stem}.pdf"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Lưu file PDF",
            str(default_path),
            "PDF Document (*.pdf)",
        )
        if not file_path:
            return

        output_path = Path(file_path)
        if output_path.suffix.lower() != ".pdf":
            output_path = output_path.with_suffix(".pdf")

        markdown_text = assistant_message.text.strip()
        document = QTextDocument()
        document.setMarkdown(markdown_text)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(str(output_path))

        try:
            document.print(printer)
        except Exception as exc:
            self._show_error(f"Không thể lưu file PDF: {exc}")
            return

        self.statusBar().showMessage(f"Đã xuất 1 phản hồi Trợ lý: {output_path.name}", 4000)
        if self.auto_open_export_checkbox is None or self.auto_open_export_checkbox.isChecked():
            self._open_exported_file(output_path)

    def _open_exported_file(self, output_path: Path) -> None:
        file_url = QUrl.fromLocalFile(str(output_path))
        opened = QDesktopServices.openUrl(file_url)
        if not opened:
            self.statusBar().showMessage(
                f"Đã lưu file nhưng không thể mở tự động: {output_path.name}",
                5000,
            )

    def _current_conversation_title(self) -> str | None:
        current_item = self.conversation_list.currentItem()
        if current_item is not None:
            maybe_title = current_item.data(Qt.ItemDataRole.UserRole + 1)
            if isinstance(maybe_title, str) and maybe_title.strip():
                return maybe_title.strip()

        if self.state.current_conversation_id:
            return f"Cuộc trò chuyện {self.state.current_conversation_id[:8]}"
        return None

    def _safe_filename(self, value: str) -> str:
        cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", value).strip().strip(".")
        return cleaned or "cuoc_tro_chuyen"

    def _show_error(self, message: str) -> None:
        self.statusBar().showMessage("Lỗi", 4000)
        QMessageBox.critical(self, "Lỗi", message)

    def _conversation_label(self, conversation: Conversation) -> str:
        if conversation.title:
            return conversation.title

        created = conversation.created_at.strftime("%Y-%m-%d %H:%M")
        return f"{created}  •  {conversation.id[:8]}..."
